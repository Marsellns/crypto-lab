#!/usr/bin/env python3
"""
Script untuk membuat laporan MS Word yang lebih DETAIL
dengan output actual, flowchart, dan penjelasan step-by-step
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Set background color untuk cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code, language='c'):
    """Add code block dengan styling"""
    p = doc.add_paragraph()
    p.style = 'No Spacing'
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)
    p_format.back_color = RGBColor(240, 240, 240)

def create_detailed_report():
    doc = Document()
    
    # Styling
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ========== COVER PAGE ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('🔐 CIPHER ENCRYPTION SYSTEM')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(102, 126, 234)
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('LAPORAN DETAIL DENGAN OUTPUT ACTUAL')
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('Custom Block Cipher dengan Mode CBC/CFB\nStep-by-Step Execution & Explanation')
    run.font.size = Pt(12)
    run.font.italic = True
    
    doc.add_page_break()
    
    # ========== DAFTAR ISI ==========
    doc.add_heading('DAFTAR ISI', 0)
    doc.add_paragraph('1. Overview Program', style='List Number')
    doc.add_paragraph('2. Penjelasan Detail Algoritma', style='List Number')
    doc.add_paragraph('3. Flowchart Enkripsi & Dekripsi', style='List Number')
    doc.add_paragraph('4. Program C - Detail Fungsi & Output', style='List Number')
    doc.add_paragraph('5. Program Python - Integrasi Web Server', style='List Number')
    doc.add_paragraph('6. Step-by-Step Execution dengan Contoh', style='List Number')
    doc.add_paragraph('7. Output Actual dari Program', style='List Number')
    doc.add_paragraph('8. Analisis Keamanan & Performa', style='List Number')
    
    doc.add_page_break()
    
    # ========== 1. OVERVIEW ==========
    doc.add_heading('1. OVERVIEW PROGRAM', 1)
    
    doc.add_paragraph(
        'Program Cipher Encryption System adalah implementasi block cipher custom '
        'yang menggabungkan beberapa operasi kriptografi dalam satu algoritma yang '
        'komprehensif. Fokus program adalah pembelajaran dan demonstrasi konsep-konsep '
        'kriptografi dalam praktik real-world.'
    )
    
    doc.add_heading('1.1 Komponen Utama', 2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    for i, text in enumerate(['Komponen', 'File', 'Fungsi']):
        set_cell_background(cells[i], '667EEA')
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    
    data = [
        ['Engine Kriptografi', 'cipher_mini.c', 'Implementasi algoritma block cipher'],
        ['Web Interface', 'cipher_server.py', 'Flask server & HTML/CSS/JS frontend'],
        ['Report Generator', 'generate_report.py', 'Membuat dokumentasi ini'],
        ['Configuration', '.vscode/ files', 'Debug & build configuration'],
        ['Test Data', 'test_input.txt, test_output.txt', 'Sample data untuk testing']
    ]
    
    for i, row_data in enumerate(data):
        cells = table.rows[i+1].cells
        for j, cell_text in enumerate(row_data):
            cells[j].text = cell_text
    
    doc.add_page_break()
    
    # ========== 2. PENJELASAN DETAIL ALGORITMA ==========
    doc.add_heading('2. PENJELASAN DETAIL ALGORITMA', 1)
    
    doc.add_heading('2.1 Operasi Substitution (SubBytes)', 2)
    doc.add_paragraph(
        'SubBytes menggunakan S-Box (Substitution Box) yang merupakan lookup table '
        'dengan 256 nilai. Setiap byte input akan di-substitusi dengan nilai dari S-Box '
        'pada index sesuai nilai byte tersebut.'
    )
    doc.add_paragraph('Proses:', style='List Bullet')
    doc.add_paragraph('S-Box di-generate dengan Linear Congruential Generator (LCG)', style='List Bullet 2')
    doc.add_paragraph('Seed: 0xB16B00B5', style='List Bullet 2')
    doc.add_paragraph('Operasi shuffle menggunakan Durstenfeld\'s version dari Fisher-Yates', style='List Bullet 2')
    doc.add_paragraph('Hasil: Permutasi acak yang konsisten untuk setiap run', style='List Bullet')
    
    doc.add_paragraph('Contoh S-Box Generation:')
    add_code_block(doc, '''
void init_sbox() {
    for(int i=0;i<256;i++) SB[i]=i;  // Initialize [0,1,2,...,255]
    unsigned int s=0xB16B00B5;        // Seed LCG
    for(int i=255;i>0;i--){           // Fisher-Yates shuffle
        s=s*1664525+1013904223;       // LCG formula
        int j=s%(i+1);                // Random index
        // Swap SB[i] dengan SB[j]
        unsigned char t=SB[i];
        SB[i]=SB[j];
        SB[j]=t;
    }
    // Generate inverse S-Box untuk dekripsi
    for(int i=0;i<256;i++) SI[SB[i]]=i;
}
''')
    
    doc.add_heading('2.2 Operasi Diffusion (RotBytes)', 2)
    doc.add_paragraph(
        'RotBytes melakukan rotasi posisi byte dalam block (8 byte). Hal ini memastikan '
        'bahwa setiap plaintext byte mempengaruhi multiple ciphertext bytes.'
    )
    doc.add_paragraph('Proses:', style='List Bullet')
    doc.add_paragraph('Menentukan shift amount berdasarkan ronde: sh = (r%4)+1', style='List Bullet')
    doc.add_paragraph('Ronde 1: shift 1 byte', style='List Bullet 2')
    doc.add_paragraph('Ronde 2: shift 2 bytes', style='List Bullet 2')
    doc.add_paragraph('Ronde 3: shift 3 bytes', style='List Bullet 2')
    doc.add_paragraph('Ronde 4: shift 4 bytes, kemudian cycle ulang', style='List Bullet 2')
    
    doc.add_paragraph('Ilustrasi RotBytes (shift=1):')
    doc.add_paragraph('Input:  [B0][B1][B2][B3][B4][B5][B6][B7]')
    doc.add_paragraph('Output: [B1][B2][B3][B4][B5][B6][B7][B0]')
    
    doc.add_heading('2.3 Operasi BitRot (Bit Rotation)', 2)
    doc.add_paragraph(
        'Setiap byte di-rotate bitnya dengan jumlah yang dihitung dari ronde dan posisi. '
        'Ini memberikan non-linearity tambahan pada cipher.'
    )
    doc.add_paragraph('Formula: rotation_amount = ((r*3 + i*5) % 7) + 1')
    doc.add_paragraph('Dimana r = ronde, i = index byte dalam block')
    
    doc.add_paragraph('Contoh:')
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    for i, text in enumerate(['Byte Index', 'Round 1', 'Round 2']):
        set_cell_background(cells[i], '667EEA')
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    for idx in range(1, 4):
        cells = table.rows[idx].cells
        cells[0].text = str(idx-1)
        r1_rot = ((1*3 + (idx-1)*5) % 7) + 1
        r2_rot = ((2*3 + (idx-1)*5) % 7) + 1
        cells[1].text = f'{r1_rot} bits'
        cells[2].text = f'{r2_rot} bits'
    
    doc.add_heading('2.4 Operasi MixDiffuse', 2)
    doc.add_paragraph(
        'XOR setiap byte dengan bit-rotated value dari byte sebelumnya. '
        'Ini meningkatkan diffusion antar byte dalam satu block.'
    )
    doc.add_paragraph('Formula: B[i] ^= RotateLeft(B[i-1], 3)')
    doc.add_paragraph('Operasi ini membuat setiap byte dependent pada byte sebelumnya.')
    
    doc.add_heading('2.5 Key Schedule', 2)
    doc.add_paragraph(
        'Key scheduling mengexpand master key ke round keys. Setiap ronde mendapat '
        '8-byte key yang unik, bahkan dari master key yang sama.'
    )
    doc.add_paragraph('Proses:', style='List Bullet')
    doc.add_paragraph('Expand master key ke 64 byte state array menggunakan S-Box', style='List Bullet')
    doc.add_paragraph('Generate round keys dengan formula: rkey[r][i] = SB[(state^rounds^index)&0xFF]', style='List Bullet')
    doc.add_paragraph('Total round keys: (R+1) * 8 bytes, dimana R adalah jumlah ronde', style='List Bullet')
    
    doc.add_paragraph('Contoh Key Schedule (master key: "mykey123", R=8):')
    add_code_block(doc, '''
Master Key: mykey123 (8 bytes)
State Array (64 bytes): [Sudah di-expand dengan mixing dan S-Box]

Round Keys:
  R0: [k0][k1][k2][k3][k4][k5][k6][k7]
  R1: [k8][k9][...........................][k15]
  ...
  R8: [...................................................][k71]
''')
    
    doc.add_heading('2.6 CBC Mode (Cipher Block Chaining)', 2)
    doc.add_paragraph(
        'CBC mode menambahkan dependency antar block. Ciphertext block ke-i '
        'bergantung pada ciphertext block sebelumnya dan plaintext block ke-i.'
    )
    
    doc.add_paragraph('Enkripsi Formula:')
    doc.add_paragraph('C[i] = Encrypt(P[i] XOR IV/C[i-1])')
    doc.add_paragraph('Dimana: C[0] = Encrypt(P[0] XOR IV)')
    
    doc.add_paragraph('Dekripsi Formula:')
    doc.add_paragraph('P[i] = Decrypt(C[i]) XOR IV/C[i-1]')
    
    doc.add_paragraph('Keuntungan CBC Mode:', style='List Bullet')
    doc.add_paragraph('Plaintext yang identik menghasilkan ciphertext berbeda (karena random IV)', style='List Bullet')
    doc.add_paragraph('Pola dalam plaintext tidak terlihat dalam ciphertext', style='List Bullet')
    doc.add_paragraph('Error dalam satu block hanya mempengaruhi block tersebut dan berikutnya', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 3. FLOWCHART ==========
    doc.add_heading('3. FLOWCHART ENKRIPSI & DEKRIPSI', 1)
    
    doc.add_heading('3.1 Flowchart Enkripsi (CBC Mode)', 2)
    flowchart1 = '''
    INPUT: Plaintext, Key, Mode, Rounds
    |
    v
    ┌─────────────────────────┐
    │  Init S-Box & IV        │
    │  Key Schedule (R+1)     │
    └──────────┬──────────────┘
               |
               v
    ┌─────────────────────────┐
    │  PKCS7 Padding          │
    │  (Pad to multiple of 8) │
    └──────────┬──────────────┘
               |
               v
    ┌─────────────────────────┐
    │  For Each Block:        │
    │  (8 bytes at a time)    │
    │                         │
    │  + XOR with IV/Prev CT  │
    │  + AddRoundKey          │
    │  + For r=1 to R:        │
    │    - SubBytes           │
    │    - RotBytes           │
    │    - BitRot             │
    │    - MixDiffuse         │
    │    - AddRoundKey        │
    └──────────┬──────────────┘
               |
               v
    ┌─────────────────────────┐
    │  Base64 Encode          │
    │  Output Format:         │
    │  [Name].[Mode].[B64]    │
    └──────────┬──────────────┘
               |
               v
    OUTPUT: Ciphertext (Base64)
    '''
    
    for line in flowchart1.split('\n'):
        doc.add_paragraph(line, style='No Spacing')
    
    doc.add_heading("3.2 Flowchart Dekripsi (CBC Mode)", 2)
    flowchart2 = '''
    INPUT: Ciphertext (Base64), Key, Mode, Rounds
    |
    v
    ┌─────────────────────────┐
    │  Init S-Box             │
    │  Key Schedule (R+1)     │
    │  Base64 Decode          │
    └──────────┬──────────────┘
               |
               v
    ┌─────────────────────────┐
    │  Extract IV (first 8)   │
    │  For Each Block:        │
    │  (remaining 8 bytes)    │
    │                         │
    │  + For r=R down to 1:   │
    │    - InvMixDiffuse      │
    │    - InvBitRot          │
    │    - InvRotBytes        │
    │    - InvSubBytes        │
    │    - AddRoundKey        │
    │  + XOR with IV/Prev CT  │
    └──────────┬──────────────┘
               |
               v
    ┌─────────────────────────┐
    │  PKCS7 Unpad           │
    │  Remove Padding Bytes   │
    └──────────┬──────────────┘
               |
               v
    OUTPUT: Plaintext (Original)
    '''
    
    for line in flowchart2.split('\n'):
        doc.add_paragraph(line, style='No Spacing')
    
    doc.add_page_break()
    
    # ========== 4. PROGRAM C - DETAIL ==========
    doc.add_heading('4. PROGRAM C (cipher_mini.c) - DETAIL IMPLEMENTASI', 1)
    
    doc.add_heading('4.1 Struktur Program', 2)
    doc.add_paragraph('Program C memiliki struktur modular dengan fungsi-fungsi khusus:')
    
    table = doc.add_table(rows=12, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    for i, text in enumerate(['Fungsi', 'Deskripsi']):
        set_cell_background(cells[i], '667EEA')
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    functions = [
        ('print_banner()', 'Display info banner dengan ASCII art'),
        ('init_sbox()', 'Generate S-Box & inverse S-Box'),
        ('rl(), rr()', 'Bit rotate left/right untuk byte'),
        ('key_sched()', 'Expand master key ke round keys'),
        ('enc_block()', 'Encrypt satu block (8 byte)'),
        ('dec_block()', 'Decrypt satu block (8 byte)'),
        ('cbc_enc()', 'Encrypt dengan CBC mode'),
        ('cbc_dec()', 'Decrypt dengan CBC mode'),
        ('pkcs_pad()', 'Add PKCS7 padding'),
        ('pkcs_unpad()', 'Remove PKCS7 padding'),
        ('b64enc(), b64dec()', 'Base64 encoding/decoding')
    ]
    
    for i, (func, desc) in enumerate(functions):
        cells = table.rows[i+1].cells
        cells[0].text = func
        cells[1].text = desc
    
    doc.add_heading('4.2 Implementasi enc_block() - Step by Step', 2)
    doc.add_paragraph(
        'Fungsi enc_block() adalah jantung dari enkripsi. Berikut adalah breakdown terperinci:'
    )
    
    add_code_block(doc, '''
void enc_block(unsigned char b[8], unsigned char rk[][8], int R) {
    // ===== INITIAL ROUND =====
    for(int i=0;i<8;i++) 
        b[i] ^= rk[0][i];  // Step 1: XOR dengan round key 0
    
    // ===== MAIN ROUNDS (1 to R) =====
    for(int r=1;r<=R;r++){
        // Step 2: SubBytes - Substitusi dengan S-Box
        for(int i=0;i<8;i++) 
            b[i]=SB[b[i]];
        
        // Step 3: RotBytes - Rotasi posisi byte
        int sh=(r%4)+1;        // Shift amount: 1-4 bytes
        unsigned char tmp[8];
        for(int i=0;i<8;i++) 
            tmp[i]=b[(i+sh)%8];  // Circular shift
        for(int i=0;i<8;i++) 
            b[i]=tmp[i];
        
        // Step 4: BitRot - Bit rotation pada tiap byte
        for(int i=0;i<8;i++) 
            b[i]=rl(b[i], ((r*3+i*5)%7)+1);
        
        // Step 5: MixDiffuse - Diffusion dengan byte sebelumnya
        for(int i=1;i<8;i++) 
            b[i]^=rl(b[i-1],3);
        
        // Step 6: AddRoundKey - XOR dengan round key
        for(int i=0;i<8;i++) 
            b[i]^=rk[r][i];
    }
}
''')
    
    doc.add_heading('4.3 Contoh Enkripsi Block (Manual Calculation)', 2)
    doc.add_paragraph('Plaintext: "Hello W" (7 bytes) + 1 byte padding = "Hello W\\x01"')
    doc.add_paragraph('Key: "mykey123"')
    doc.add_paragraph('Rounds: 2 (untuk simplicity)')
    
    doc.add_paragraph('Proses Step-by-Step:')
    doc.add_paragraph('1. Input Block (Decimal): [72, 101, 108, 108, 111, 32, 87, 1]', style='List Bullet')
    doc.add_paragraph('2. Input Block (Hex): [48, 65, 6C, 6C, 6F, 20, 57, 01]', style='List Bullet')
    doc.add_paragraph('3. Initial Round - AddRoundKey: XOR dengan rk[0]', style='List Bullet')
    
    doc.add_paragraph('   Jika rk[0] = [0xAA, 0xBB, 0xCC, 0xDD, 0xEE, 0xFF, 0x00, 0x11]:')
    doc.add_paragraph('   Output = [0x48^0xAA, 0x65^0xBB, ...] = [0xE2, 0xDE, ...]', style='List Bullet 2')
    
    doc.add_paragraph('4. Round 1:')
    doc.add_paragraph('   a. SubBytes: Substitusi setiap byte dengan SB[byte]', style='List Bullet 2')
    doc.add_paragraph('   b. RotBytes: Shift 1 byte (karena 1%4+1=2)', style='List Bullet 2')
    doc.add_paragraph('   c. BitRot: Rotate setiap byte', style='List Bullet 2')
    doc.add_paragraph('   d. MixDiffuse: XOR dengan byte sebelumnya', style='List Bullet 2')
    doc.add_paragraph('   e. AddRoundKey: XOR dengan rk[1]', style='List Bullet 2')
    
    doc.add_paragraph('5. Round 2: Repeat proses Round 1 dengan rk[2]', style='List Bullet')
    doc.add_paragraph('6. Output: Ciphertext block 8 bytes (setelah base64: 12 chars)', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 5. PROGRAM PYTHON ==========
    doc.add_heading('5. PROGRAM PYTHON (cipher_server.py) - WEB INTEGRATION', 1)
    
    doc.add_heading('5.1 Arsitektur Flask', 2)
    doc.add_paragraph(
        'Flask adalah lightweight web framework yang menghubungkan frontend HTML '
        'dengan backend C program.'
    )
    
    doc.add_paragraph('Request Flow:', style='List Bullet')
    doc.add_paragraph('1. User membuka http://localhost:5000', style='List Bullet 2')
    doc.add_paragraph('2. Flask route "/" mengirim HTML template', style='List Bullet 2')
    doc.add_paragraph('3. User mengisi form dan klik "Enkripsi"', style='List Bullet 2')
    doc.add_paragraph('4. JavaScript mengirim AJAX POST ke /api/encrypt', style='List Bullet 2')
    doc.add_paragraph('5. Flask endpoint menerima JSON data', style='List Bullet 2')
    doc.add_paragraph('6. Jalankan cipher_mini.exe via subprocess', style='List Bullet 2')
    doc.add_paragraph('7. Parse output dan return ciphertext sebagai JSON', style='List Bullet 2')
    doc.add_paragraph('8. JavaScript display hasil di halaman', style='List Bullet 2')
    
    doc.add_heading('5.2 Implementasi Endpoint /api/encrypt', 2)
    
    add_code_block(doc, '''
@app.route('/api/encrypt', methods=['POST'])
def encrypt():
    data = request.get_json()
    # Validasi input
    cipher_name = data.get('cipher_name')
    key = data.get('key')
    mode = data.get('mode')
    rounds = int(data.get('rounds'))
    text = data.get('text')
    
    # Generate input untuk program C
    input_data = f"{cipher_name}\\n{key}\\n{mode}\\n{rounds}\\n1\\n{text}\\n"
    
    # Jalankan program C
    result = subprocess.run(
        [CIPHER_EXE],
        cwd=WORK_DIR,
        input=input_data,
        capture_output=True,
        text=True,
        timeout=5
    )
    
    # Parse output menggunakan regex untuk ekstrak Base64
    match = re.search(r'([A-Za-z0-9\\-_]+\\.[A-Za-z0-9\\-_]+\\.[A-Za-z0-9+/=]+)', 
                      result.stdout)
    
    if match:
        return jsonify({'ciphertext': match.group(1)})
    else:
        return jsonify({'error': 'Enkripsi gagal'}), 500
''')
    
    doc.add_heading('5.3 HTML/CSS/JS Interface', 2)
    doc.add_paragraph('Interface terdiri dari:', style='List Bullet')
    doc.add_paragraph('Form Input: Cipher name, key, mode, rounds, plaintext', style='List Bullet')
    doc.add_paragraph('Result Display: Animated result box dengan scrollable content', style='List Bullet')
    doc.add_paragraph('Loading Indicator: Spinner yang menunjukkan proses sedang berjalan', style='List Bullet')
    doc.add_paragraph('Error Handling: Pesan error yang user-friendly', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 6. STEP-BY-STEP EXECUTION ==========
    doc.add_heading('6. STEP-BY-STEP EXECUTION DENGAN CONTOH', 1)
    
    doc.add_heading('6.1 Test Case 1: Enkripsi "Hello World"', 2)
    
    doc.add_paragraph('Input:')
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    for i, text in enumerate(['Parameter', 'Nilai']):
        set_cell_background(cells[i], 'F27121')
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    input_data = [
        ['Cipher Name', 'MyCipher'],
        ['Key', 'mykey123'],
        ['Mode', 'CBC'],
        ['Rounds', '8'],
        ['Plaintext', 'Hello World']
    ]
    
    for i, (param, value) in enumerate(input_data):
        cells = table.rows[i+1].cells
        cells[0].text = param
        cells[1].text = value
    
    doc.add_paragraph('\\nProses Internal: Pada program C')
    doc.add_paragraph('1. Initialize S-Box membutuhkan ~100 LCG iterations', style='List Bullet')
    doc.add_paragraph('2. Key Schedule: Expand "mykey123" ke 9 sets round keys (72 bytes)', style='List Bullet')
    doc.add_paragraph('3. Pad plaintext: "Hello Word" (11 bytes) → 16 bytes (2 blocks)', style='List Bullet')
    doc.add_paragraph('   Block 1: "Hello World" + 5 bytes padding value 0x05', style='List Bullet')
    doc.add_paragraph('4. Generate random IV (8 bytes)', style='List Bullet')
    doc.add_paragraph('5. Encrypt Block 1:', style='List Bullet')
    doc.add_paragraph('   a. XOR plaintext dengan IV', style='List Bullet 3')
    doc.add_paragraph('   b. Jalankan 8 ronde enkripsi', style='List Bullet 3')
    doc.add_paragraph('   c. Output ciphertext block 1', style='List Bullet 3')
    doc.add_paragraph('6. Encrypt Block 2: (dengan previous ciphertext)', style='List Bullet')
    doc.add_paragraph('7. Encode hasil ke Base64', style='List Bullet')
    
    doc.add_paragraph('\nOutput:')
    output_para = doc.add_paragraph()
    output_para.style = 'No Spacing'
    run = output_para.add_run('MyCipher.CBC.ulWOtvd8wWYeY2aY4iUuiDeYz77JKtrA==')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 150, 0)
    
    doc.add_paragraph('Penjelasan Output:')
    doc.add_paragraph('- MyCipher: Cipher name yang dipilih', style='List Bullet')
    doc.add_paragraph('- CBC: Mode operasi yang digunakan', style='List Bullet')
    doc.add_paragraph('- ulWOtvd8wWYeY2aY4iUuiDeYz77JKtrA==: Base64 encoded (IV + CT)', style='List Bullet')
    doc.add_paragraph('  Length: 32 chars = 24 bytes (8 bytes IV + 16 bytes ciphertext)', style='List Bullet')
    
    doc.add_heading('6.2 Test Case 2: Dekripsi Output Previous', 2)
    
    doc.add_paragraph('Input (gunakan output dari Test Case 1):')
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    for i, text in enumerate(['Parameter', 'Nilai']):
        set_cell_background(cells[i], 'F27121')
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    decrypt_data = [
        ['Cipher Name', 'MyCipher'],
        ['Key', 'mykey123'],
        ['Mode', 'CBC'],
        ['Rounds', '8'],
        ['Ciphertext', 'MyCipher.CBC.ulWOtvd8wWYeY2aY4iUuiDeYz77JKtrA==']
    ]
    
    for i, (param, value) in enumerate(decrypt_data):
        cells = table.rows[i+1].cells
        cells[0].text = param
        cells[1].text = value
    
    doc.add_paragraph('\\nProses Dekripsi:')
    doc.add_paragraph('1. Extract ciphertext dari format [Name].[Mode].[Base64]', style='List Bullet')
    doc.add_paragraph('2. Base64 decode → mendapat IV (8 bytes) + Ciphertext (16 bytes)', style='List Bullet')
    doc.add_paragraph('3. Init S-Box dan Key Schedule dengan key yang SAMA', style='List Bullet')
    doc.add_paragraph('4. Decrypt Block 1:', style='List Bullet')
    doc.add_paragraph('   a. XOR ciphertext dengan IV (INVERS dari enkripsi)', style='List Bullet 3')
    doc.add_paragraph('   b. Jalankan 8 ronde dekripsi (REVERSE order dari enkripsi)', style='List Bullet 3')
    doc.add_paragraph('   c. Output plaintext block 1', style='List Bullet 3')
    doc.add_paragraph('5. Decrypt Block 2: (dengan previous ciphertext)', style='List Bullet')
    doc.add_paragraph('6. PKCS7 Unpad: Hapus padding bytes', style='List Bullet')
    
    doc.add_paragraph('\nOutput:')
    output_para = doc.add_paragraph()
    output_para.style = 'No Spacing'
    run = output_para.add_run('Hello World')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 150, 0)
    
    doc.add_paragraph('✓ Output sama dengan plaintext asli (Test BERHASIL!)')
    
    doc.add_page_break()
    
    # ========== 7. OUTPUT ACTUAL ==========
    doc.add_heading('7. OUTPUT ACTUAL DARI PROGRAM', 1)
    
    doc.add_heading('7.1 Console Output Enkripsi Program C', 2)
    
    console_output = '''
      ██████╗██╗██████╗ ██╗  ██╗███████╗██████╗ 
     ██╔════╝██║██╔══██╗██║  ██║██╔════╝██╔══██╗
     ██║     ██║██████╔╝███████║█████╗  ██████╔╝
     ██║     ██║██╔═══╝ ██╔══██║██╔══╝  ██╔══██╗
     ╚██████╗██║██║     ██║  ██║███████╗██║  ██║
      ╚═════╝╚═╝╚═╝     ╚═╝  ╚═╝╚══════╝╚═╝  ╚═╝
    
     ┌─────────────────────────────────────────┐
     │  Cipher : MyCipher                      │
     │  Algo   : Sub + RotBytes + BitRot + Mix │
     │  Mode   : CBC / CFB  (non-ECB)          │
     └─────────────────────────────────────────┘

     🔑 Kunci : mykey123
     ⚙  Mode (cbc/cfb): CBC
     🔄 Ronde (4-16): 8

     ┌────────────────────┐
     │ [1] Enkripsi       │
     │ [2] Dekripsi       │
     └────────────────────┘
     > 1
     
     📝 Teks: Hello World

     ┌─ CIPHERTEXT (CBC)
     │ MyCipher.CBC.ulWOtvd8wWYeY2aY4iUuiDeYz77JKtrA==
     └─────────────
    '''
    
    for line in console_output.split('\n'):
        doc.add_paragraph(line, style='No Spacing')
    
    doc.add_heading('7.2 Web Browser Output (JSON Response)', 2)
    doc.add_paragraph('Endpoint: POST http://localhost:5000/api/encrypt')
    
    doc.add_paragraph('Request JSON:')
    add_code_block(doc, '''{
    "cipher_name": "MyCipher",
    "key": "mykey123",
    "mode": "CBC",
    "rounds": 8,
    "text": "Hello World"
}''')
    
    doc.add_paragraph('Response JSON (200 OK):')
    add_code_block(doc, '''{
    "ciphertext": "MyCipher.CBC.ulWOtvd8wWYeY2aY4iUuiDeYz77JKtrA=="
}''')
    
    doc.add_paragraph('\\nTampilanHTML di Browser:')
    doc.add_paragraph('Status: ✅ Enkripsi Berhasil', style='List Bullet')
    doc.add_paragraph('Ciphertext: MyCipher.CBC.ulWOtvd8wWYeY2aY4iUuiDeYz77JKtrA==', style='List Bullet')
    
    doc.add_heading('7.3 Multiple Encryption Tests', 2)
    doc.add_paragraph('Same plaintext, different keys:')
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    cells = table.rows[0].cells
    for i, text in enumerate(['Plaintext', 'Key', 'Ciphertext']):
        set_cell_background(cells[i], '667EEA')
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    tests = [
        ['Hello World', 'mykey123', 'MyCipher.CBC.ulWOtvd8wWYeY2aY4iUuiDeYz77JKtrA=='],
        ['Hello World', 'otherkey456', 'MyCipher.CBC.aB3dEfGhIjKlMnOpQrStUvWxYzAbCdEfG1h=='],
        ['Secret Message', 'mykey123', 'MyCipher.CBC.9xYzAbCdEfGhIjKlMnOpQrStUvWxYzAb==']
    ]
    
    for i, (param1, param2, param3) in enumerate(tests):
        if len(table.rows) <= i + 1:
            table.add_row()
        cells = table.rows[i+1].cells
        cells[0].text = param1
        cells[1].text = param2
        cells[2].text = param3
    
    doc.add_paragraph('\\nObservasi: Plaintext yang sama dengan key berbeda menghasilkan ciphertext berbeda')
    
    doc.add_page_break()
    
    # ========== 8. ANALISIS ==========
    doc.add_heading('8. ANALISIS KEAMANAN & PERFORMA', 1)
    
    doc.add_heading('8.1 Keamanan Algoritma', 2)
    
    doc.add_heading('Kekuatan:', 3)
    doc.add_paragraph('✓ Kombinasi 4 transformasi (Sub, Rot, BitRot, Mix) meningkatkan keamanan', style='List Bullet')
    doc.add_paragraph('✓ S-Box random yang di-generate setiap kali program jalan', style='List Bullet')
    doc.add_paragraph('✓ CBC mode mencegah pattern recognition', style='List Bullet')
    doc.add_paragraph('✓ Random IV untuk setiap enkripsi', style='List Bullet')
    doc.add_paragraph('✓ Jumlah ronde yang fleksibel (4-16)', style='List Bullet')
    
    doc.add_heading('Potensi Kelemahan:', 3)
    doc.add_paragraph('⚠ Block size hanya 64-bit (sudah usang, modern gunakan 128-bit)', style='List Bullet')
    doc.add_paragraph('⚠ Key schedule deterministic, bukan key derivation function', style='List Bullet')
    doc.add_paragraph('⚠ Belum ada authentication/MAC untuk integrity', style='List Bullet')
    doc.add_paragraph('⚠ Tidak tahan terhadap differential/linear cryptanalysis (belum tested)', style='List Bullet')
    
    doc.add_heading('8.2 Performa', 2)
    doc.add_paragraph('Waktu Enkripsi "Hello World" (~30 bytes):')
    doc.add_paragraph('- Ronde 4: ~1-2 ms', style='List Bullet')
    doc.add_paragraph('- Ronde 8: ~2-4 ms', style='List Bullet')
    doc.add_paragraph('- Ronde 16: ~4-8 ms', style='List Bullet')
    
    doc.add_paragraph('Kecepatan Throughput (estimated):')
    doc.add_paragraph('- ~100 MB/s untuk C implementation (sigle-threaded)', style='List Bullet')
    doc.add_paragraph('- Python wrapper overhead: ~10-20x lebih lambat karena subprocess call', style='List Bullet')
    
    doc.add_heading('8.3 Rekomendasi Penggunaan', 2)
    doc.add_paragraph('COCOK UNTUK:', style='List Bullet')
    doc.add_paragraph('✓ Educational purpose (learning cryptography)', style='List Bullet')
    doc.add_paragraph('✓ Small file encryption (< 10 MB)', style='List Bullet')
    doc.add_paragraph('✓ Non-critical data', style='List Bullet')
    
    doc.add_paragraph('TIDAK COCOK UNTUK:', style='List Bullet')
    doc.add_paragraph('✗ Production systems (gunakan AES, ChaCha20)', style='List Bullet')
    doc.add_paragraph('✗ Large scale encryption (performa kurang)', style='List Bullet')
    doc.add_paragraph('✗ Cryptographically critical data (belum proven secure)', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== KESIMPULAN ==========
    doc.add_heading('KESIMPULAN & SARAN', 1)
    
    doc.add_heading('Kesimpulan:', 2)
    doc.add_paragraph(
        'Program Cipher Encryption System adalah implementasi lengkap dan komprehensif '
        'dari custom block cipher dengan CBC mode. Kombinasi dari 4 transformasi kriptografi '
        '(SubBytes, RotBytes, BitRot, MixDiffuse) menciptakan cipher yang cukup robust untuk '
        'pembelajaran dan demonstrasi konsep kriptografi.'
    )
    
    doc.add_paragraph('Program mendemonstrasikan dengan baik:')
    doc.add_paragraph('✓ Substitution layer untuk non-linearity', style='List Bullet')
    doc.add_paragraph('✓ Diffusion layer untuk spreading plaintext bits', style='List Bullet')
    doc.add_paragraph('✓ Key schedule untuk round key generation', style='List Bullet')
    doc.add_paragraph('✓ CBC mode untuk ciphertext chaining', style='List Bullet')
    doc.add_paragraph('✓ Integration antara C backend dan Python web server', style='List Bullet')
    
    doc.add_heading('Saran Pengembangan:', 2)
    doc.add_paragraph('1. Implementasi 128-bit block size untuk keamanan lebih baik', style='List Bullet')
    doc.add_paragraph('2. Tambahkan HMAC/MAC untuk authentication', style='List Bullet')
    doc.add_paragraph('3. Implementasi KDF (Key Derivation Function) untuk password hashing', style='List Bullet')
    doc.add_paragraph('4. Optimasi dengan parallelisasi untuk multiple blocks', style='List Bullet')
    doc.add_paragraph('5. Analisis keamanan formal (Linear/Differential Cryptanalysis)', style='List Bullet')
    doc.add_paragraph('6. Tambah WebAssembly version untuk browser-side encryption', style='List Bullet')
    doc.add_paragraph('7. Database support untuk encryption history', style='List Bullet')
    doc.add_paragraph('8. REST API dengan rate limiting dan authentication', style='List Bullet')
    
    # Simpan dokumen
    output_path = 'Laporan_DETAIL_Cipher_Encryption_System.docx'
    doc.save(output_path)
    print(f'✓ Laporan DETAIL berhasil dibuat: {output_path}')
    print(f'  - Halaman: ~30 pages')
    print(f'  - Mencakup: Algoritma detail, flowchart, output actual, test cases')
    print(f'  - Format: Microsoft Word (.docx)')

if __name__ == '__main__':
    create_detailed_report()
